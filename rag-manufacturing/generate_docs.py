"""生成示例工业文档（docx + pdf 双格式，含目录和页码）"""
import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from fpdf import FPDF

DOCX_DIR = "data/docx"
PDF_DIR = "data/pdf"


# ============================================================
#  工具函数：docx 目录 + 页码
# ============================================================
def add_toc(doc):
    """插入目录域（Word打开后右键更新域即可自动生成目录）"""
    doc.add_heading("目录", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'begin'); run._r.append(fld)
    run2 = p.add_run()
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'; run2._r.append(instr)
    run3 = p.add_run()
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate'); run3._r.append(fld2)
    run4 = p.add_run('[打开Word后，右键此处 → 更新域，即可自动生成目录]')
    run4.font.size = Pt(10)
    run5 = p.add_run()
    fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end'); run5._r.append(fld3)
    doc.add_page_break()


def add_page_number_footer(doc):
    """给所有节添加页码：第X页/共Y页"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = 1  # 居中
        p.add_run("第 ")

        r1 = p.add_run()
        f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); r1._r.append(f1)
        r2 = p.add_run(); i1 = OxmlElement('w:instrText'); i1.text = 'PAGE'; r2._r.append(i1)
        r3 = p.add_run()
        f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate'); r3._r.append(f2)
        r4 = p.add_run('1')
        r5 = p.add_run()
        f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end'); r5._r.append(f3)

        p.add_run(" 页 / 共 ")

        r6 = p.add_run()
        f4 = OxmlElement('w:fldChar'); f4.set(qn('w:fldCharType'), 'begin'); r6._r.append(f4)
        r7 = p.add_run(); i2 = OxmlElement('w:instrText'); i2.text = 'NUMPAGES'; r7._r.append(i2)
        r8 = p.add_run()
        f5 = OxmlElement('w:fldChar'); f5.set(qn('w:fldCharType'), 'separate'); r8._r.append(f5)
        r9 = p.add_run('1')
        r10 = p.add_run()
        f6 = OxmlElement('w:fldChar'); f6.set(qn('w:fldCharType'), 'end'); r10._r.append(f6)

        p.add_run(" 页")


# ============================================================
#  PDF 基类：带页码
# ============================================================
class NumberedPDF(FPDF):
    def footer(self):
        self.set_y(-15)
        self.set_font("st", size=9)
        self.cell(0, 10, f"第 {self.page_no()} 页", align="C")

    def draw_table(self, headers: list, rows: list, col_widths: list = None):
        """在PDF中绘制表格
        headers: 表头列表
        rows: 数据行列表，每行也是一个列表
        col_widths: 每列宽度，None则均分
        """
        if col_widths is None:
            usable = self.w - 2 * self.l_margin
            col_widths = [usable / len(headers)] * len(headers)
        line_h = 7  # 行高

        # ── 表头 ──
        self.set_font("st", "B", 10)
        self.set_fill_color(220, 220, 220)
        for i, (h, w) in enumerate(zip(headers, col_widths)):
            self.cell(w, line_h, h, border=1, fill=True, align="C")
        self.ln()

        # ── 数据行 ──
        self.set_font("st", size=9)
        for row_idx, row in enumerate(rows):
            # 交替行背景色
            if row_idx % 2 == 0:
                self.set_fill_color(250, 250, 250)
            else:
                self.set_fill_color(255, 255, 255)

            # 计算本行最大高度
            max_lines = 1
            cell_texts = []
            for i, (cell_text, w) in enumerate(zip(row, col_widths)):
                text = str(cell_text)
                # 估算需要的行数
                char_per_line = max(1, int(w / self.get_string_width("测")))
                lines = (len(text) + char_per_line - 1) // char_per_line
                max_lines = max(max_lines, lines)
                cell_texts.append((text, lines))

            row_h = line_h * max_lines
            x_start = self.get_x()

            for i, (cell_text, lines) in enumerate(cell_texts):
                w = col_widths[i]
                self.set_xy(x_start + sum(col_widths[:i]), self.get_y())
                # 用 multi_cell 画多行文本
                self.multi_cell(w, line_h, cell_text, border=1, fill=True, align="L")

            self.set_xy(x_start, self.get_y())

        self.ln(4)

# ============================================================
#  文档1：注塑机操作手册
# ============================================================
def generate_manual():
    # ---------- DOCX ----------
    doc = Document()
    add_toc(doc)  # 插入目录（Word打开后更新域即可）
    doc.add_heading("注塑机操作手册", level=0)

    doc.add_heading("第一章 设备概述", level=1)
    doc.add_heading("1.1 设备简介", level=2)
    doc.add_paragraph("HTF1200型卧式注塑机是本公司主力生产设备，采用伺服电机驱动和"
                       "双缸注射系统，适用于热塑性塑料的精密成型加工。最大锁模力1200吨，"
                       "注射量2000~8000克，可加工PP、ABS、PA、PC等常用工程塑料。")

    doc.add_heading("1.2 主要技术参数", level=2)
    t = doc.add_table(rows=8, cols=3, style="Light Grid Accent 1")
    for r, (k, v, u) in enumerate([
        ("螺杆直径", "80/90/100", "mm"), ("最大锁模力", "1200", "吨"),
        ("注射压力", "180", "MPa"), ("注射速率", "450", "cm³/s"),
        ("螺杆转速", "0-120", "rpm"), ("加热功率", "45", "kW"),
        ("工作电压", "380", "V/50Hz"),
    ]):
        t.cell(r + 1, 0).text = k
        t.cell(r + 1, 1).text = v
        t.cell(r + 1, 2).text = u

    doc.add_heading("第二章 操作流程", level=1)
    doc.add_heading("2.1 开机前检查", level=2)
    doc.add_paragraph("1. 检查液压油位是否在油标中位以上。\n"
                       "2. 确认冷却水阀门已打开，水压在0.3-0.5MPa。\n"
                       "3. 检查安全门行程开关是否灵敏。\n"
                       "4. 确认料筒内无异物，料斗有料。\n"
                       "5. 检查各润滑点是否有足够的润滑油。")

    doc.add_heading("2.2 正常启动步骤", level=2)
    doc.add_paragraph("1. 合上主电源开关，控制面板亮起。\n"
                       "2. 启动液压泵电机，检查运转方向是否正确。\n"
                       "3. 设置料筒各段加热温度，等待升温至设定值（约30分钟）。\n"
                       "4. 温度达到后，手动低速对空注射，确认熔料均匀无气泡。\n"
                       "5. 安装模具，调整锁模力至工艺卡要求。\n"
                       "6. 手动模式下试模3-5次，确认无异常后切换至自动模式。")

    doc.add_heading("2.3 停机操作", level=2)
    doc.add_paragraph("1. 将料筒温度降至保温温度（150-180度）。\n"
                       "2. 关闭料斗下料口，排空料筒余料。\n"
                       "3. 模具合模但不加压，关闭液压泵。\n"
                       "4. 断开主电源，做好交接记录。")

    doc.add_heading("第三章 参数设置", level=1)
    doc.add_heading("3.1 温度参数", level=2)
    t = doc.add_table(rows=6, cols=3, style="Light Grid Accent 1")
    for r, (z, temp, dev) in enumerate([
        ("第一段（料斗侧）", "180-200", "±5"), ("第二段", "200-220", "±5"),
        ("第三段", "220-240", "±5"), ("第四段", "230-250", "±5"),
        ("喷嘴", "240-260", "±3"),
    ]):
        t.cell(r + 1, 0).text = z; t.cell(r + 1, 1).text = temp; t.cell(r + 1, 2).text = dev

    doc.add_heading("3.2 压力参数", level=2)
    t = doc.add_table(rows=5, cols=3, style="Light Grid Accent 1")
    for r, (p, v, n) in enumerate([
        ("注射压力", "80-180", "根据材料和模具调整"),
        ("保压压力", "60-120", "一般为注射压力的60%-80%"),
        ("背压", "5-15", "螺杆回退时料筒内压力"),
        ("锁模压力", "800-1200", "根据模具面积计算"),
    ]):
        t.cell(r + 1, 0).text = p; t.cell(r + 1, 1).text = v; t.cell(r + 1, 2).text = n

    doc.add_heading("第四章 安全规范", level=1)
    doc.add_heading("4.1 个人防护要求", level=2)
    doc.add_paragraph("操作人员必须穿戴防烫手套、安全鞋和护目镜。禁止佩戴手套操作控制面板。")

    doc.add_heading("4.2 紧急停止程序", level=2)
    doc.add_paragraph("发生以下情况立即按下急停按钮（红色蘑菇头）：\n"
                       "1. 模具或机械部件出现异常响声或振动。\n"
                       "2. 料筒温度失控持续上升。\n"
                       "3. 液压管路爆裂或严重泄漏。\n"
                       "4. 人员或异物进入合模区域。\n"
                       "急停后需由班组长和维修人员共同确认故障排除后方可复位。")

    doc.add_heading("4.3 禁止操作清单", level=2)
    t = doc.add_table(rows=6, cols=2, style="Light Grid Accent 1")
    for r, (item, reason) in enumerate([
        ("禁止屏蔽安全门开关", "可能导致人员重伤或死亡"),
        ("禁止在自动模式下手动操作", "PLC程序冲突可能导致设备损坏"),
        ("禁止超过最大注射量生产", "导致螺杆过载断裂"),
        ("禁止在无冷却水情况下加热", "导致料筒变形报废"),
        ("禁止擅自修改PLC参数", "导致生产工艺失控"),
    ]):
        t.cell(r + 1, 0).text = item; t.cell(r + 1, 1).text = reason

    add_page_number_footer(doc)
    doc.save(os.path.join(DOCX_DIR, "01_注塑机操作手册.docx"))

    # ---------- PDF ----------
    pdf = NumberedPDF()
    pdf.add_font("st", "", "c:/windows/fonts/simsun.ttc")
    pdf.add_font("st", "B", "c:/windows/fonts/simhei.ttf")
    pdf.set_auto_page_break(True, 15)
    pdf.add_page()
    pdf.set_font("st", "B", 18)
    pdf.cell(0, 12, "注塑机操作手册", align="C")
    pdf.ln(14)

    # Ch1
    pdf.set_font("st", "B", 14)
    pdf.start_section("第一章 设备概述", level=1)
    pdf.cell(0, 10, "第一章 设备概述"); pdf.ln(10)
    pdf.set_font("st", "B", 12)
    pdf.start_section("1.1 设备简介", level=2)
    pdf.cell(0, 8, "1.1 设备简介"); pdf.ln(8)
    pdf.set_font("st", size=11)
    pdf.multi_cell(0, 6.5, "HTF1200型卧式注塑机是本公司主力生产设备，采用伺服电机驱动和双缸注射系统，"
                            "适用于热塑性塑料的精密成型加工。最大锁模力1200吨，注射量2000~8000克，"
                            "可加工PP、ABS、PA、PC等常用工程塑料。")
    pdf.ln(4)
    pdf.set_font("st", "B", 12)
    pdf.start_section("1.2 主要技术参数", level=2)
    pdf.cell(0, 8, "1.2 主要技术参数"); pdf.ln(8)
    pdf.draw_table(
        ["参数", "数值", "单位"],
        [["螺杆直径", "80/90/100", "mm"], ["最大锁模力", "1200", "吨"],
         ["注射压力", "180", "MPa"], ["注射速率", "450", "cm³/s"],
         ["螺杆转速", "0-120", "rpm"], ["加热功率", "45", "kW"],
         ["工作电压", "380", "V/50Hz"]],
        [55, 35, 30]
    )

    # Ch2
    pdf.set_font("st", "B", 14)
    pdf.start_section("第二章 操作流程", level=1)
    pdf.cell(0, 10, "第二章 操作流程"); pdf.ln(10)
    pdf.set_font("st", "B", 12)
    pdf.start_section("2.1 开机前检查", level=2)
    pdf.cell(0, 8, "2.1 开机前检查"); pdf.ln(8)
    pdf.set_font("st", size=11)
    for s in ["1. 检查液压油位是否在油标中位以上。",
              "2. 确认冷却水阀门已打开，水压在0.3-0.5MPa。",
              "3. 检查安全门行程开关是否灵敏。",
              "4. 确认料筒内无异物，料斗有料。",
              "5. 检查各润滑点是否有足够的润滑油。"]:
        pdf.cell(0, 6.5, s); pdf.ln(6.5)
    pdf.ln(4)
    pdf.set_font("st", "B", 12)
    pdf.start_section("2.2 正常启动步骤", level=2)
    pdf.cell(0, 8, "2.2 正常启动步骤"); pdf.ln(8)
    pdf.set_font("st", size=11)
    for s in ["1. 合上主电源开关，控制面板亮起。",
              "2. 启动液压泵电机，检查运转方向是否正确。",
              "3. 设置料筒各段加热温度，等待升温至设定值（约30分钟）。",
              "4. 温度达到后，手动低速对空注射，确认熔料均匀无气泡。",
              "5. 安装模具，调整锁模力至工艺卡要求。",
              "6. 手动模式下试模3-5次，确认无异常后切换至自动模式。"]:
        pdf.cell(0, 6.5, s); pdf.ln(6.5)
    pdf.ln(4)
    pdf.set_font("st", "B", 12)
    pdf.start_section("2.3 停机操作", level=2)
    pdf.cell(0, 8, "2.3 停机操作"); pdf.ln(8)
    pdf.set_font("st", size=11)
    for s in ["1. 将料筒温度降至保温温度（150-180度）。",
              "2. 关闭料斗下料口，排空料筒余料。",
              "3. 模具合模但不加压，关闭液压泵。",
              "4. 断开主电源，做好交接记录。"]:
        pdf.cell(0, 6.5, s); pdf.ln(6.5)

    # Ch3
    pdf.set_font("st", "B", 14)
    pdf.start_section("第三章 参数设置", level=1)
    pdf.cell(0, 10, "第三章 参数设置"); pdf.ln(10)
    pdf.set_font("st", "B", 12)
    pdf.start_section("3.1 温度参数", level=2)
    pdf.cell(0, 8, "3.1 温度参数"); pdf.ln(8)
    pdf.draw_table(
        ["加热区段", "设定温度(度)", "允许偏差(度)"],
        [["第一段(料斗侧)", "180-200", "±5"], ["第二段", "200-220", "±5"],
         ["第三段", "220-240", "±5"], ["第四段", "230-250", "±5"],
         ["喷嘴", "240-260", "±3"]],
        [50, 40, 35]
    )
    pdf.ln(2)
    pdf.set_font("st", "B", 12)
    pdf.start_section("3.2 压力参数", level=2)
    pdf.cell(0, 8, "3.2 压力参数"); pdf.ln(8)
    pdf.draw_table(
        ["参数", "数值(MPa)", "说明"],
        [["注射压力", "80-180", "根据材料和模具调整"],
         ["保压压力", "60-120", "一般为注射压力的60%-80%"],
         ["背压", "5-15", "螺杆回退时料筒内压力"],
         ["锁模压力", "800-1200", "根据模具面积计算"]],
        [40, 35, 55]
    )

    # Ch4
    pdf.set_font("st", "B", 14)
    pdf.start_section("第四章 安全规范", level=1)
    pdf.cell(0, 10, "第四章 安全规范"); pdf.ln(10)
    pdf.set_font("st", "B", 12)
    pdf.start_section("4.1 个人防护要求", level=2)
    pdf.cell(0, 8, "4.1 个人防护要求"); pdf.ln(8)
    pdf.set_font("st", size=11)
    pdf.multi_cell(0, 6.5, "操作人员必须穿戴防烫手套、安全鞋和护目镜。禁止佩戴手套操作控制面板。")
    pdf.ln(4)
    pdf.set_font("st", "B", 12)
    pdf.start_section("4.2 紧急停止程序", level=2)
    pdf.cell(0, 8, "4.2 紧急停止程序"); pdf.ln(8)
    pdf.set_font("st", size=11)
    pdf.multi_cell(0, 6.5, "发生以下情况立即按下急停按钮：1.模具或机械部件出现异常响声或振动。"
                            "2.料筒温度失控持续上升。3.液压管路爆裂或严重泄漏。"
                            "4.人员或异物进入合模区域。急停后需由班组长和维修人员共同确认故障排除后方可复位。")
    pdf.ln(4)
    pdf.set_font("st", "B", 12)
    pdf.start_section("4.3 禁止操作清单", level=2)
    pdf.cell(0, 8, "4.3 禁止操作清单"); pdf.ln(8)
    pdf.draw_table(
        ["禁止操作", "后果"],
        [["屏蔽安全门开关", "可能导致人员重伤或死亡"],
         ["自动模式下手动操作", "PLC程序冲突可能导致设备损坏"],
         ["超过最大注射量生产", "导致螺杆过载断裂"],
         ["无冷却水情况下加热", "导致料筒变形报废"],
         ["擅自修改PLC参数", "导致生产工艺失控"]],
        [60, 60]
    )
    pdf.output(os.path.join(PDF_DIR, "01_注塑机操作手册.pdf"))
    print("  01_注塑机操作手册.docx / .pdf 生成完毕")


# ============================================================
#  文档2：故障码手册
# ============================================================
def generate_fault_codes():
    E_CODES = [
        ("E01", "加热器断路", "加热圈电阻值异常（正常15-25Ω）",
         "1.万用表测量加热圈电阻 2.更换断路加热圈 3.检查固态继电器"),
        ("E02", "热电偶故障", "温度显示----或跳变",
         "1.检查热电偶接线端子是否松动 2.用标准温度计比对 3.更换热电偶"),
        ("E03", "电机过载", "电机电流超过额定值120%持续3秒",
         "1.检查机械传动是否卡滞 2.测量三相电压平衡 3.降低负载后重启"),
        ("E04", "变频器故障", "变频器面板显示Err代码",
         "1.记录Err代码查阅变频器手册 2.断电5分钟后重启 3.联系厂家技术支持"),
        ("E05", "安全门开关异常", "前后安全门行程开关信号不一致",
         "1.检查安全门是否完全关闭 2.清洁行程开关触点 3.更换损坏的行程开关"),
        ("E06", "接近开关失效", "PLC未检测到到位信号",
         "1.检查接近开关指示灯 2.调整感应距离(正常3-5mm) 3.更换接近开关"),
        ("E07", "电磁阀卡死", "阀芯不动作或动作缓慢",
         "1.拆洗阀芯及阀套 2.检查液压油清洁度 3.更换磨损的电磁阀"),
        ("E08", "编码器信号丢失", "位置显示异常或归零",
         "1.检查编码器联轴器是否松动 2.测量供电电压(DC5V/24V) 3.更换编码器"),
    ]
    AL_CODES = [
        ("AL-01", "料筒温度偏高", "实际温度>设定值+15度",
         "1.检查冷却风扇是否运转 2.降低该段加热功率 3.检查热电偶是否紧贴料筒"),
        ("AL-02", "液压油温过高", "油温>55度",
         "1.检查冷却水流量 2.清洗冷却器水垢 3.降低系统工作压力"),
        ("AL-03", "润滑油不足", "润滑泵压力<0.5MPa",
         "1.补充润滑油 2.检查润滑管路是否堵塞 3.检查润滑泵电机"),
        ("AL-04", "料斗缺料", "光电传感器检测不到原料",
         "1.及时补充原料 2.检查光电传感器镜片是否被灰尘遮挡 3.调整传感器灵敏度"),
        ("AL-05", "模具温度异常", "模温机显示温度偏差>10度",
         "1.检查模温机循环水 2.检查模具冷却水道是否堵塞 3.维修模温机"),
        ("AL-06", "注射时间过长", "实际注射时间>设定值+2秒",
         "1.检查喷嘴是否堵塞 2.提高注射速度 3.升高料筒温度"),
        ("AL-07", "储料不足", "螺杆回退位置未达到设定值",
         "1.检查料斗下料是否顺畅 2.增加背压 3.检查止逆环是否磨损"),
        ("AL-08", "液压油位低", "油箱液位低于下限",
         "1.补充46号抗磨液压油 2.检查是否有泄漏点 3.修复泄漏后补充至中位"),
    ]
    P_CODES = [
        ("P01", "注射压力不足", "实际注射压力<设定值-10MPa",
         "1.检查比例压力阀线圈 2.清洗比例阀 3.检查液压泵磨损情况"),
        ("P02", "锁模力不足", "锁模油缸压力<设定值",
         "1.检查锁模油缸密封圈 2.调整锁模压力补偿 3.检查液压管路泄漏"),
        ("P03", "保压不稳定", "保压阶段压力波动>5MPa",
         "1.检查蓄能器氮气压力 2.检查比例阀响应 3.PID参数优化"),
        ("P04", "背压异常", "背压无法建立或过高",
         "1.检查背压调节阀 2.清洁液压先导阀 3.校准压力传感器"),
        ("P05", "射胶压力冲击", "注射起始阶段压力峰值过高",
         "1.调整注射加速度曲线 2.检查止逆环密封 3.降低注射起始速度"),
    ]

    # ---------- DOCX ----------
    doc = Document()
    add_toc(doc)
    doc.add_heading("注塑机故障码手册", level=0)
    doc.add_paragraph("本手册收录HTF1200型注塑机常见故障码，包含电气系统故障（E系列）、"
                       "报警信息（AL系列）及压力系统故障（P系列）三大类，共21条。")

    for title, codes in [("第一章 电气系统故障（E系列）", E_CODES),
                          ("第二章 报警信息（AL系列）", AL_CODES),
                          ("第三章 压力系统故障（P系列）", P_CODES)]:
        doc.add_heading(title, level=1)
        for code, name, cause, solution in codes:
            doc.add_heading(f"{code} - {name}", level=2)
            t = doc.add_table(rows=2, cols=2, style="Light Grid Accent 1")
            t.cell(0, 0).text = "可能原因"; t.cell(0, 1).text = cause
            t.cell(1, 0).text = "解决方案"; t.cell(1, 1).text = solution

    add_page_number_footer(doc)
    doc.save(os.path.join(DOCX_DIR, "02_故障码手册.docx"))

    # ---------- PDF ----------
    pdf = NumberedPDF()
    pdf.add_font("st", "", "c:/windows/fonts/simsun.ttc")
    pdf.add_font("st", "B", "c:/windows/fonts/simhei.ttf")
    pdf.set_auto_page_break(True, 15)
    pdf.add_page()
    pdf.set_font("st", "B", 18)
    pdf.cell(0, 12, "注塑机故障码手册", align="C")
    pdf.ln(14)

    for title, codes in [("第一章 电气系统故障（E系列）", E_CODES),
                          ("第二章 报警信息（AL系列）", AL_CODES),
                          ("第三章 压力系统故障（P系列）", P_CODES)]:
        pdf.set_font("st", "B", 14)
        pdf.start_section(title, level=1)
        pdf.cell(0, 10, title); pdf.ln(10)
        for code, name, cause, solution in codes:
            pdf.set_font("st", "B", 12)
            pdf.start_section(f"{code} - {name}", level=2)
            pdf.cell(0, 8, f"{code} - {name}"); pdf.ln(8)
            pdf.draw_table(
                ["项目", "内容"],
                [["可能原因", cause], ["解决方案", solution]],
                [25, 95]
            )
            pdf.ln(2)

    pdf.output(os.path.join(PDF_DIR, "02_故障码手册.pdf"))
    print("  02_故障码手册.docx / .pdf 生成完毕")


# ============================================================
#  文档3：维保操作规范
# ============================================================
def generate_maintenance():
    DAILY = [
        ("检查液压油位", "目视", "油标中位以上", "补充46号抗磨液压油"),
        ("检查润滑油量", "目视", "油位在上限与下限之间", "补充导轨润滑油"),
        ("清洁模具表面", "手动", "无残留塑料、无油污", "使用铜刷和抹布"),
        ("检查安全门功能", "功能测试", "开关门时机器立即停止或允许启动", "报修"),
        ("检查冷却水", "目视", "水压0.3-0.5MPa，无泄漏", "调整阀门、修复泄漏"),
        ("清理料斗磁力架", "手动", "无金属杂物", "清理并记录异物来源"),
    ]
    WEEKLY = [
        ("检查各紧固螺栓", "扳手", "无松动", "按扭矩要求紧固"),
        ("清洁液压油冷却器", "手动", "翅片无灰尘堵塞", "压缩空气吹扫"),
        ("检查皮带张紧度", "张力计", "挠度8-12mm", "调整电机底座"),
        ("润滑拉杆及导轨", "手动", "表面均匀油膜", "涂抹润滑脂"),
        ("检查电气接线端子", "目视+螺丝刀", "无松动、无烧焦痕迹", "紧固或更换"),
    ]
    MONTHLY = [
        ("更换液压油滤芯", "更换", "滤芯使用<=500小时", "更换新滤芯"),
        ("检查螺杆磨损", "测量", "螺杆与料筒间隙<=0.3mm", "超差则更换螺杆"),
        ("校准温度传感器", "标准温度计比对", "偏差<=正负2度", "修正或更换热电偶"),
        ("检查止逆环", "拆卸检查", "密封面光滑无凹槽", "磨损则更换"),
        ("测试紧急停止", "功能测试", "按下急停后所有动作立即停止", "检查急停回路"),
    ]
    QUARTERLY = [
        ("更换液压油", "更换", "油质清亮无乳化", "更换46号抗磨液压油，清洗油箱"),
        ("检查电机绝缘", "兆欧表", "绝缘电阻>=10MΩ", "烘干或更换电机"),
        ("校验压力传感器", "标准压力表比对", "偏差<=1%FS", "校准或更换传感器"),
        ("清洗液压阀组", "拆卸清洗", "阀芯运动灵活无卡滞", "超声波清洗后复装"),
        ("PLC系统备份", "程序备份", "备份成功校验通过", "导出程序并归档"),
    ]

    # ---------- DOCX ----------
    doc = Document()
    add_toc(doc)
    doc.add_heading("注塑机维保操作规范", level=0)

    for section_title, items in [("第一章 日常维保（每班次）", DAILY),
                                  ("第二章 周度维保", WEEKLY),
                                  ("第三章 月度维保", MONTHLY),
                                  ("第四章 季度维保", QUARTERLY)]:
        doc.add_heading(section_title, level=1)
        t = doc.add_table(rows=len(items) + 1, cols=4, style="Light Grid Accent 1")
        for c, h in enumerate(["检查项目", "方法", "合格标准", "异常处理"]):
            t.cell(0, c).text = h
        for r, (item, method, std, action) in enumerate(items):
            t.cell(r + 1, 0).text = item
            t.cell(r + 1, 1).text = method
            t.cell(r + 1, 2).text = std
            t.cell(r + 1, 3).text = action

    add_page_number_footer(doc)
    doc.save(os.path.join(DOCX_DIR, "03_维保操作规范.docx"))

    # ---------- PDF ----------
    pdf = NumberedPDF()
    pdf.add_font("st", "", "c:/windows/fonts/simsun.ttc")
    pdf.add_font("st", "B", "c:/windows/fonts/simhei.ttf")
    pdf.set_auto_page_break(True, 15)
    pdf.add_page()
    pdf.set_font("st", "B", 18)
    pdf.cell(0, 12, "注塑机维保操作规范", align="C")
    pdf.ln(14)

    for section_title, items in [("第一章 日常维保（每班次）", DAILY),
                                  ("第二章 周度维保", WEEKLY),
                                  ("第三章 月度维保", MONTHLY),
                                  ("第四章 季度维保", QUARTERLY)]:
        pdf.set_font("st", "B", 14)
        pdf.start_section(section_title, level=1)
        pdf.cell(0, 10, section_title); pdf.ln(10)
        pdf.draw_table(
            ["检查项目", "方法", "合格标准", "异常处理"],
            [[item, method, std, action] for item, method, std, action in items],
            [35, 22, 38, 30]
        )
        pdf.ln(4)

    pdf.output(os.path.join(PDF_DIR, "03_维保操作规范.pdf"))
    print("  03_维保操作规范.docx / .pdf 生成完毕")


if __name__ == "__main__":
    os.makedirs(DOCX_DIR, exist_ok=True)
    os.makedirs(PDF_DIR, exist_ok=True)
    print("开始生成工业文档...\n")
    generate_manual()
    generate_fault_codes()
    generate_maintenance()
    print(f"\n生成完毕！共6个文件 -> {DOCX_DIR}/ 和 {PDF_DIR}/")
