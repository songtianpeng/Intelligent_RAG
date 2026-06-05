# 第一步 先定metadata Schema
"""
首先确定好一个chunk和metadata的Schema结构
{
    "page_content": "E03 电机过载指的是电机电流超过额定值120%...",
    "metadata": {
        "source": "02_故障码手册.pdf",          # 来源文件名
        "doc_title": "注塑机故障码手册",         # 文档标题（PDF/Word里的主标题）
        "page": 17,                             # 第几页（PDF有，DOCX统一为1）
        "heading_path": "第一章 电气系统故障 > E03 电机过载",  # WIKI式层级路径
        "chunk_type": "table"                   # text / table / list
    }
}
"""


# 第二步，写骨架
"""
文档处理管道 — PageIndex + WIKI式层级切分
每个chunk带来源文件名、页码、章节路径、内容类型
"""
import os
import re # 可引用正则表达式，验证邮箱或者某种特殊格式
from typing import List, Optional, Tuple  # 给变量、参数、返回值增加类型说明，提高代码可读性、IDE提示能力和静态检查能力

import fitz # fitz 指的是 PyMuPDF库 安装命令 pip install pymupdf
from docx import Document as DocxDocument  # 用来读取word文档，指定别名
from langchain_core.documents import Document  # 导入的是 LangChain 的 文档对象（Document Object）
from langchain_text_splitters import RecursiveCharacterTextSplitter # 递归字符分割器
from lxml.etree import QName # 用来识别word类型是文本还是表格


class DocumentLoader:
    def __init__(self, data_dir: str = "../data/"):
        self.data_dir = data_dir
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=500, chunk_overlap=50,
            separators=["\n\n", "\n", "。", "，", " "]
        )

    # ── 公开入口 ──
    def load_all(self) -> List[Document]:
        all_docs = []
        for root, _, files in os.walk(self.data_dir):   # os.walk() 是什么，递归遍历指定目录及其所有子目录 root 当前目录 files 当前目录下的所有文件
            for f in sorted(files):  # sorted() 函数 对文件进行一个排序再遍历
                path = os.path.join(root, f) # 拼接路径
                if f.endswith(".docx"): # 如果是docx,word文档
                    all_docs.extend(self._load_docx(path))  # 调用 _load_docx() 加载方法
                elif f.endswith(".pdf"): # 如果是pdf文档
                    all_docs.extend(self._load_pdf(path)) # 调用 _load_pdf() 加载方法
        for doc in all_docs:
            doc.page_content = doc.metadata["heading_path"] + "\n" + doc.page_content
        print(f"[DocumentLoader] 加载了 {len(all_docs)} 页/段") # 打印加载的文档数量
        chunks = self._split_large_pages(all_docs) # 调用 _split_large_pages() 切分、分割方法
        print(f"[DocumentLoader] 切分后 {len(chunks)} chunks") # 打印切分后的块数量
        return chunks

    def _load_pdf(self, path: str) -> List[Document]:
        """
        加载PDF文档
        :param path: 文件全路径名
        :return: 存储了文档对象列表，包含每一页文档内容、元数据
        """
        source = os.path.basename(path)  # 从完整文件路径中提取文件名。不带路径
        pdf = fitz.open(path) # 打开PDF文件
        doc_title = self._extract_doc_title_from_first_page(pdf) # 提取PDF文档的标题

        docs = [] # 存储文档对象列表
        heading_stack = []  # 当前层级路径栈
        current_text_parts = []  # 当前 heading 下的正文累积
        current_heading_path = ""  # 当前 heading 的完整路径

        for page_num in range(len(pdf)):
            # ★ 新增：检测并提取表格
            tables = pdf[page_num].find_tables()
            if tables:
                for table in tables:
                    rows_data = table.extract()  # 返回 List[List[str]]
                    if rows_data and len(rows_data) > 0:
                        lines = ['\t'.join(str(c or '') for c in row) for row in rows_data]
                        table_text = '\n'.join(lines)
                        heading_path = source
                        docs.append(self._make_doc(
                            table_text, source, doc_title,
                            page=page_num + 1, heading_path=heading_path,
                            chunk_type="table"
                        ))
            text = pdf[page_num].get_text()
            if not text.strip(): # 如果文本为空，则跳过，text.strip() 返回一个字符串
                continue
            for line in text.split("\n"):
                line = line.strip()
                if not line: # 如果文本为空，则跳过
                    continue
                level = self._get_heading_level_pdf(line) # 获取标题层级
                if level: # 如果是标题
                    # ── 遇到新标题，先提交上一段累积的正文 ──
                    if current_text_parts:  # 如果有累积的正文
                        docs.append(self._make_doc(
                            "\n".join(current_text_parts), source, doc_title,
                            page=page_num + 1, heading_path=current_heading_path,
                            chunk_type=self._detect_chunk_type("\n".join(current_text_parts))
                        ))
                        current_text_parts = []

                    # ── 更新层级栈 ──
                    # level=1 → 清空所有；level=2 → 保留 level=1，替换 level=2
                    heading_stack = heading_stack[:level - 1]  # 删除多余的层级，每个层级的标题只保留对应的路径
                    heading_stack.append(line)  # 一级标题，只保留了一级标题，在上一步已清空所有
                    current_heading_path = " > ".join(heading_stack)  # 构建层级路径
                else:
                    # ── 正文：累积到当前 heading 下 ──
                    current_text_parts.append(line)  # 不是标题就累积正文
        # 最后一段
        if current_text_parts:  # 因为最后一段，没有新标题了，如果累积正文有值，直接提交，正文和来源都在前面已存储好
            # current_text_parts.append(current_heading_path)  # 将标题路径也添加到chunk内容里面，提高语义检索和关键词检索的准确率
            docs.append(self._make_doc(
                "\n".join(current_text_parts), source, doc_title,
                page=len(pdf), heading_path=current_heading_path,
                chunk_type=self._detect_chunk_type("\n".join(current_text_parts))
            ))
        pdf.close() # 关闭PDF文件，释放资源
        return docs

    def _load_docx(self, path: str) -> List[Document]:
        """
        docx文件加载
        :param path: 文档全路径名
        :return: 存储了文档对象列表，包含文档内容、元数据
        """
        source = os.path.basename(path) # 从完整文件路径中提取文件名
        doc = DocxDocument(path) # 打开docx文件
        doc_title = source.replace(".docx", "")  # 文件名即文档标题

        docs = []
        heading_stack = []           # 当前层级路径栈
        current_text_parts = []      # 当前 heading 下的正文累积
        current_heading_path = ""    # 当前 heading 的完整路径

        # 给 body 里每个子元素打上序号
        body_children = list(doc.element.body)
        position = {el: i for i, el in enumerate(body_children)}
        # 段落 + 表格混在一起，按文档顺序排
        all_items = []
        for para in doc.paragraphs:
            all_items.append((position.get(para._element, 9999), 'p', para))
        for table in doc.tables:
            all_items.append((position.get(table._element, 9999), 'tbl', table))
        all_items.sort(key=lambda x: x[0])

        # 遍历时直接用 python-docx 的原生对象
        for pos, item_type, obj in all_items:
            if item_type == 'p':
                # obj 是 Paragraph → .style.name / .text 全都能用
                text = obj.text.strip()
                if not text:
                    continue

                level = self._get_heading_level(obj)  # 获取标题层级
                if level is not None:  # 如果是标题
                    # ── 遇到新标题，先提交上一段累积的正文 ──
                    if current_text_parts:  # 如果有累积的正文
                        # current_text_parts.append(current_heading_path) # 将标题路径也添加到chunk内容里面，提高语义检索和关键词检索的准确率
                        docs.append(self._make_doc(
                            "\n".join(current_text_parts), source, doc_title,
                            page=1, heading_path=current_heading_path,
                            chunk_type=self._detect_chunk_type("\n".join(current_text_parts))
                        ))
                        current_text_parts = []

                    # ── 更新层级栈 ──
                    # level=1 → 清空所有；level=2 → 保留 level=1，替换 level=2
                    heading_stack = heading_stack[:level - 1]  # 删除多余的层级，每个层级的标题只保留对应的路径
                    heading_stack.append(text)  # 一级标题，只保留了一级标题，在上一步已清空所有
                    current_heading_path = " > ".join(heading_stack)  # 构建层级路径
                else:
                    # ── 正文：累积到当前 heading 下 ──
                    current_text_parts.append(text)  # 不是标题就累积正文
            elif item_type == 'tbl':
                # obj 是 Table → .rows / .cells / .cell(row, col).text 全都能用
                table_text = self._get_table_text(obj)  # 获取表格内容
                current_text_parts.append(table_text)  # 将表格同样添加到正文累积
        # 最后一段
        if current_text_parts: # 因为最后一段，没有新标题了，如果累积正文有值，直接提交，正文和来源都在前面已存储好
            # current_text_parts.append(current_heading_path)  # 将标题路径也添加到chunk内容里面，提高语义检索和关键词检索的准确率
            docs.append(self._make_doc(
                "\n".join(current_text_parts), source, doc_title,
                page=1, heading_path=current_heading_path,
                chunk_type=self._detect_chunk_type("\n".join(current_text_parts))
            ))
        return docs

    def _get_heading_level(self, para) -> Optional[int]:
        """
        判断标题层级以及是否是标题
        :param para: 传进来的段落内容
        :return: 标题的层级，None代表不是标题
        """
        """返回 1/2/3 如果该段是标题，否则 None"""
        style_name = para.style.name if para.style else "" # 获取段落样式名称
        # Word 内置样式: "Heading 1" / "Heading 2" / "Heading 3"
        # 标题层级覆盖到四级，基本上已经占到所有文档的95%以上了，足够知道回答溯源到哪一个章节了，后面不在细化，会保留在正文里
        for level in range(1, 5): # 遍历标题层级，左闭右开
            if f"Heading {level}" in style_name:
                return level
            # # 也有可能是中文名 标题 1，标题 2，标题 3，正文
            if f"标题 {level}" in style_name:
                return level
        # 正则兜底：文本模式匹配
        text = para.text.strip()
        # 正则兜底
        patterns = [
            (4, r"^\d+\.\d+\.\d+\.\d+"),
            (3, r"^\d+\.\d+\.\d+"),
            (2, r"^\d+\.\d+"),
            (1, r"^第[一二三四五六七八九十百千万\d]+章"),
            (1, r"^\d+(?!\.)(?:\s*)[^\d\s].*$"),
            (1, r"^附录[A-Z]"),
        ]
        # 每一个和匹配的标题层级模板，都进行正则匹配
        for level, pattern in patterns:
            if re.match(pattern, text): # 如果匹配成功，返回标题层级
                return level
        return None # 如果不是标题，返回 None，代表是正文

    def _get_heading_level_pdf(self, text: str) -> Optional[int]:
        """
        判断标题层级以及是否是标题
        :param para: 传进来的一行内容
        :return: 标题的层级，None代表不是标题
        """
        """返回 1/2/3 如果该段是标题，否则 None"""
        # 正则兜底：文本模式匹配
        text = text.strip()
        # 正则兜底
        patterns = [
            (4, r"^\d+\.\d+\.\d+\.\d+"),
            (3, r"^\d+\.\d+\.\d+"),
            (2, r"^\d+\.\d+"),
            (1, r"^第[一二三四五六七八九十百千万\d]+章"),
            (1, r"^\d+(?!\.)(?:\s*)[^\d\s].*$"),
            (1, r"^附录[A-Z]"),
        ]
        # 每一个和匹配的标题层级模板，都进行正则匹配
        for level, pattern in patterns:
            if re.match(pattern, text): # 如果匹配成功，返回标题层级
                return level
        return None # 如果不是标题，返回 None，代表是正文

    def _split_large_pages(self, docs: List[Document]) -> List[Document]:
        """
        切分页内容或者chunk内容。<=500字 的页直接保留，>500字 的页在子标题处切。表格页永远不切。
        :param docs: 页内容或chunk内容
        :return: 切分更小的chunk的内容
        """
        chunks = []
        for doc in docs:
            content = doc.page_content
            # 表格不受切分
            if doc.metadata.get("chunk_type") == "table":
                chunks.append(doc)
                continue
            # 小页直接保留
            if len(content) <= 500:
                chunks.append(doc)
                continue
            # 大页：在子标题处切
            sub_parts = self._split_by_sub_heading(content) # 已按照子标题切分出更小块chunk
            for part_text, part_heading in sub_parts:    # 遍历更小块chunk,两个参数分别是正文和标题
                new_heading = doc.metadata["heading_path"]
                if part_heading:
                    new_heading = f"{new_heading} > {part_heading}" if new_heading else part_heading # 构建新的标题路径
                chunks.append(Document(
                    page_content=part_text,
                    metadata={**doc.metadata, "heading_path": new_heading} # 解包有相同的键，后面覆盖前面的值
                ))
        return chunks

    def _split_by_sub_heading(self, text: str) -> List[Tuple[str, str]]:
        """在子标题处切分，返回 [(文本, 子标题), ...]"""
        pattern = re.compile(
            r"^(\d+(?:\.\d+){1,3})\s+(.+)",
            re.MULTILINE
        )
        parts = [] # 存放最终结果
        last_pos = 0 # 上一个子标题的结束位置
        last_heading = "" # 上一个子标题

        for m in pattern.finditer(text): # 扫描所有标题，并遍历
            if m.start() > last_pos: # 如果当前标题的开始位置大于上一个标题的结束位置，说明中间有内容，需要切分
                parts.append((text[last_pos:m.start()].strip(), last_heading)) # 添加上一个标题的文本和标题元组
            # 判断完是否要切分，其实当前标题的任务已经完成了，当前标题变成了上一个标题
            last_heading = m.group(0).strip() # 更新上一个标题
            last_pos = m.start() # 更新上一个标题的结束位置

        if last_pos < len(text): # 如果当前位置小于文本长度，说明有剩余内容，需要切分，即处理最后一个标题未处理内容
            parts.append((text[last_pos:].strip(), last_heading)) # 添加剩余内容和最后一个标题

        return parts or [(text, "")] # 如果没有匹配到标题，返回一个空标题

    def _make_doc(
        self, content: str, source: str, doc_title: str,
        page: int, heading_path: str, chunk_type: str
    ) -> Document:
        """
        构建Document对象
        :param content: 文本内容
        :param source: 文件路径
        :param doc_title: 文档标题
        :param page: 页数
        :param heading_path: 溯源标题层级
        :param chunk_type: 类型，文本，表格...
        :return:
        """
        return Document(
            page_content=content,
            metadata={
                "source": source,
                "doc_title": doc_title,
                "page": page,
                "heading_path": heading_path,
                "chunk_type": chunk_type,
            }
        )

    def _extract_doc_title_from_first_page(self, pdf: fitz.Document) -> str:
        """从PDF第一页提取文档标题"""
        if len(pdf) == 0:
            return "未知文档"
        text = pdf[0].get_text()
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        if not lines:
            return "未知文档"

        # 标题特征：长度适中（6-40字）、不是纯数字、不含列表符号
        candidates = []
        for line in lines[:5]:  # 只看前5行
            if 6 <= len(line) <= 40 and not re.match(r"^\d+$", line):
                candidates.append(line)

        return candidates[0] if candidates else lines[0]

    def _detect_chunk_type(self, text: str) -> str:
        """检测chunk内容类型：text / table / list"""
        lines = [l for l in text.split("\n") if l.strip()]
        total_lines = len(lines)
        if total_lines == 0:
            return "text"

        # 表格特征：含大量制表符 \t 或多个空格分隔的列结构
        tab_count = sum(1 for l in lines if "\t" in l)
        # 文本中用 3个以上连续空格模拟的列结构
        column_count = sum(1 for l in lines if "   " in l)

        if tab_count >= 2 or column_count >= total_lines * 0.5:
            return "table"

        # 列表特征：30%以上的行以数字或符号开头
        list_pattern = re.compile(
            r"^\s*(?:"
            r"\d+[\.\)、．]|"  # 1. 1) 1）
            r"[一二三四五六七八九十]+、|"  # 一、二、三、
            r"[①②③④⑤⑥⑦⑧⑨⑩]|"  # ①②③
            r"（\d+）"  # （1）（2）
            r")"
        )
        list_starts = sum(1 for l in lines if list_pattern.match(l))
        if list_starts >= total_lines * 0.3:
            return "list"

        return "text"

    def _get_table_text(self, element) -> str:
        """
        获取表格内容
        :param element: 表格
        :return: 表格内容
        """
        table_content = []
        for row in element.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            row_text = "\t".join(row_data)
            table_content.append(row_text)
        table_text = "\n".join(table_content)
        return table_text



if __name__ == '__main__':
    # pass
    loader = DocumentLoader()
    chunks = loader.load_all()
    for i,chunk in enumerate(chunks):
        print(f"第{i+1}chunk,内容是：{chunk.page_content},溯源信息是：{chunk.metadata}")

    # source = os.path.basename("../data/docx/01_注塑机操作手册.docx")  # 从完整文件路径中提取文件名
    # doc = DocxDocument("../data/docx/01_注塑机操作手册.docx")  # 打开docx文件
    # doc_title = source.replace(".docx", "")  # 文件名即文档标题
    #
    # docs = []
    # heading_stack = []  # 当前层级路径栈
    # current_text_parts = []  # 当前 heading 下的正文累积
    # current_heading_path = ""  # 当前 heading 的完整路径
    #
    # for element in doc.element.body:
    #     print(element)
    #
    # print("===================================")
    # for idx, table in enumerate(doc.tables):
    #     print(f"\n===== 表格 {idx + 1} =====")
    #
    #     for row in table.rows:
    #         row_data = [cell.text.strip() for cell in row.cells]
    #         print(row_data)
    #
    # print("====================================")
    # for paragraph in doc.paragraphs:
    #     print(paragraph.text)
    # def load_pdf(self, path: str) -> List[Document]:
    #     source = os.path.basename(path)
    #     pdf = fitz.open(path)
    #     doc_title = self._extract_doc_title_from_first_page(pdf)
    #
    #     docs = []
    #     for page_num in range(len(pdf)):
    #         page = pdf[page_num]
    #         # ★ 新增：检测并提取表格
    #         tables = page.find_tables()
    #         if tables:
    #             for table in tables:
    #                 rows_data = table.extract()  # 返回 List[List[str]]
    #                 if rows_data and len(rows_data) > 0:
    #                     lines = ['\t'.join(str(c or '') for c in row) for row in rows_data]
    #                     table_text = '\n'.join(lines)
    #                     print(table_text)
    #                     heading_path = source
    #                     docs.append(self._make_doc(
    #                         table_text, source, doc_title,
    #                         page=page_num + 1, heading_path = heading_path,
    #                         chunk_type="table"
    #                     ))
    #
    #         # 整页文本（保持原有逻辑）
    #         text = page.get_text()
    #         if not text.strip():
    #             continue
    #         heading_path = source
    #         docs.append(Document(
    #             page_content=text.strip(),
    #             metadata={
    #                 "source": source,
    #                 "doc_title": doc_title,
    #                 "page": page_num + 1,
    #                 "heading_path": heading_path,
    #                 "chunk_type": self._detect_chunk_type(text),
    #             }
    #         ))
    #     pdf.close()
    #     return docs


    # docs = load_pdf(loader, "../data/pdf/01_注塑机操作手册.pdf")
    # for doc in docs:
    #     print(doc.page_content, doc.metadata)